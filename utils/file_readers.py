import os
from typing import List, Literal, Union
from pdf2image import convert_from_path
import pytesseract
from tqdm import tqdm

from langchain_community.document_loaders import (
    PyMuPDFLoader,
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredPDFLoader,
)
from langchain_core.documents import Document
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def read_pdf_file(
    filepath: str,
    pdfloader: Literal["PYMUPDF", "PYPDF", "unstructured-langchain", "TESSERACT"] = "PYMUPDF",
    mode: Literal["single", "multi", "all"] = "single",
) -> List[Document]:
    """
    Reads a PDF file and extracts its content using the specified PDF loader.

    Args:
        filepath (str): The path to the PDF file.
        pdfloader (Literal): The PDF loader to use. Options are "PYMUPDF", "PYPDF", "unstructured-langchain", "TESSERACT".
        mode (Literal, optional): The mode used for unstructured-langchain. Options are "single", "multi" or "all".

    Returns:
        List[Document]: A list of LangChain `Document` objects containing extracted text and metadata.
    
    Raises:
        ValueError: If an unsupported `pdfloader` is provided.
    """
    if pdfloader == "PYMUPDF":
        loader = PyMuPDFLoader(filepath)
    elif pdfloader == "PYPDF":
        loader = PyPDFLoader(filepath)
    elif pdfloader == "unstructured-langchain":
        loader = UnstructuredPDFLoader(filepath, mode=mode)
    elif pdfloader == "TESSERACT":
        images = convert_from_path(filepath, thread_count=24, dpi=300, grayscale=True)
        items = [
            Document(
                page_content=pytesseract.image_to_string(image, lang="en"),
                metadata={"file_path": filepath, "page": page_no, "total_pages": len(images)},
            )
            for page_no, image in tqdm(enumerate(images), desc="Processing PDF pages")
        ]
        return items
    else:
        raise ValueError("PDF Loader Not Supported")
    return loader.load()


def extract_table_text(table: Table) -> str:
    """
    Extracts text content from a table in a DOCX document.

    Args:
        table (Table): A `Table` object from `python-docx`.

    Returns:
        str: A string representation of the table, with rows separated by newlines and columns separated by tabs.
    """
    return "\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows)


def read_docx_file(
    filepath: str,
    docxloader: Literal["python-docx", "docx2txt", "unstructured-langchain"] = "python-docx",
    mode: Literal["single", "multi", "all"] = "single",
) -> List[Document]:
    """
    Reads a DOCX file and extracts its content using the specified DOCX loader.

    Args:
        filepath (str): The path to the DOCX file.
        docxloader (Literal): The DOCX loader to use. Options are "python-docx", "docx2txt", "unstructured-langchain".
        mode (Literal): The mode used for unstructured-langchain. Options are "single", "multi" or "all".

    Returns:
        List[Document]: A list of LangChain `Document` objects containing extracted text and metadata.

    Raises:
        ValueError: If an unsupported `docxloader` is provided.
    """
    if docxloader == "unstructured-langchain":
        loader = UnstructuredWordDocumentLoader(filepath, mode=mode)
        return loader.load()
    
    if docxloader == "docx2txt":
        loader = Docx2txtLoader(filepath)
        return loader.load()
    
    if docxloader == "python-docx":
        doc = DocxDocument(filepath)
        items = []
        _, filename = os.path.split(filepath)
        chunk_no = 1

        for element in doc.element.body:
            meta = {"file_path": filepath, "filename": filename}

            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                if para.text.strip():
                    all_bold = all(run.bold for run in para.runs if run.text.strip())
                    all_underline = all(run.underline for run in para.runs if run.text.strip())

                    meta["element_style"] = "bold" if all_bold else "underline" if all_underline else "paragraph"
                    meta["center_align"] = para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                    meta["chunk_no"] = chunk_no

                    items.append(Document(page_content=para.text, metadata=meta))
                    chunk_no += 1

            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                meta.update({"element_style": "table", "center_align": False, "chunk_no": chunk_no})

                items.append(Document(page_content=extract_table_text(table), metadata=meta))
                chunk_no += 1

        return items

    raise ValueError("DOCX Loader Not Supported")
